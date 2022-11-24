from operator import truediv
import docx
# import aspose.words as aw

# doc = aw.Document("arquivo.pdf")
# doc.save("out.docx")

doc = docx.Document('out.docx')

lista_cidade = [
   'S.J.RPRETOÀBADYBASSITTR$5,30TempoMédiodoPercurso:50minutos'
   'BADYBASSITTÀSJRIOPRETO',
   'SJRIOPRETOÀNOVAALIANÇAR$6,30TempoMédiodoPercurso:1Hora',
   'NOVAALIANÇAÀSJRIOPRETO',
   'SJRIOPRETOÀNOVAITAPIREMAR$7,50TempoMédiodoPercurso:1Horae15minutos',
   'NOVAITAPIREMAÀSJRIOPRETO',
   'SJRIOPRETOÀMENDONÇAR$10,35TempoMédiodoPercurso:1Horae30minutos',
   'MENDONÇAÀSJRIOPRETO',
   'SJRIOPRETOÀADOLFOR$12,25TempoMédiodoPercurso:1Horae45minutos',
   'ADOLFOÀSJRIOPRETO',
   'SJRIOPRETOÀMIRASSOLÂNDIAR$6,55TempoMédiodoPercurso:1Hora',
   'MIRASSOLÂNDIAÀSJRIOPRETO',
   'SJRIOPRETOÀIPIGUÁR$5,30TempoMédiodoPercurso:45minutos',
   'IPIGUÁÀSJRIOPRETO',
   'SJRIOPRETOÀRECANTO18R$5,30TempoMédiodoPercurso:20minutos',
   'RECANTO18ÀSJRIOPRETO',
   'SJRIOPRETOÀMACAÚBASR$7,50TempoMédiodoPercurso:1Horae15Minutos',
   'MACAÚBASÀSJRIOPRETO',
   
]

lista_dias = [
   'SEGUNDAASEXTA-FEIRA',
   'SEGUNDAASEXTA-FEIRA',
   'DOMINGOSEFERIADOS',
   'DOMINGOEFERIADOS',
   'SÁBADO',
   'SÁBADOS',
]

# Variaveis 
index = 0
cidade = False
sabado = False
domingo = False
id = 999


for i in doc.paragraphs:


   linha= " ".join((i.text).strip().replace(" ","").split())

   if linha != '':
      if linha in lista_cidade:
         if linha == lista_cidade[0]:
            cidade = True
            id = 0


   else:
      pass



